"""
Convert paper.md to a formatted Word document using python-docx.
Handles headings, tables, code blocks, bold/italic, and bullet lists.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import re


def create_word_from_md():
    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # ── Default style ──
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    # Set East Asian font
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = doc.styles['Normal'].element.get_or_add_rPr().makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')

    # ── Parse markdown ──
    lines = []
    with open('paper.md', 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Skip horizontal rules
        if line.strip() == '---':
            i += 1
            continue

        # H1 -> Title
        if line.startswith('# ') and not line.startswith('## '):
            title_text = line[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(title_text)
            run.bold = True
            run.font.size = Pt(18)
            run.font.name = '黑体'
            rPr = run._element.get_or_add_rPr()
            rFonts_elem = rPr.makeelement(qn('w:rFonts'), {})
            rFonts_elem.set(qn('w:eastAsia'), '黑体')
            rPr.insert(0, rFonts_elem)
            i += 1
            continue

        # H2 -> Heading 1
        if line.startswith('## '):
            text = line[3:].strip()
            p = doc.add_heading(text, level=1)
            for run in p.runs:
                run.font.name = '黑体'
                run.font.size = Pt(16)
                rPr = run._element.get_or_add_rPr()
                rFonts_elem = rPr.makeelement(qn('w:rFonts'), {})
                rFonts_elem.set(qn('w:eastAsia'), '黑体')
                rPr.insert(0, rFonts_elem)
            i += 1
            continue

        # H3 -> Heading 2
        if line.startswith('### '):
            text = line[4:].strip()
            p = doc.add_heading(text, level=2)
            for run in p.runs:
                run.font.name = '黑体'
                run.font.size = Pt(14)
                rPr = run._element.get_or_add_rPr()
                rFonts_elem = rPr.makeelement(qn('w:rFonts'), {})
                rFonts_elem.set(qn('w:eastAsia'), '黑体')
                rPr.insert(0, rFonts_elem)
            i += 1
            continue

        # H4 -> Heading 3
        if line.startswith('#### '):
            text = line[5:].strip()
            p = doc.add_heading(text, level=3)
            for run in p.runs:
                run.font.name = '黑体'
                run.font.size = Pt(13)
                rPr = run._element.get_or_add_rPr()
                rFonts_elem = rPr.makeelement(qn('w:rFonts'), {})
                rFonts_elem.set(qn('w:eastAsia'), '黑体')
                rPr.insert(0, rFonts_elem)
            i += 1
            continue

        # Code block
        if line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].rstrip().startswith('```'):
                code_lines.append(lines[i].rstrip())
                i += 1
            i += 1  # skip closing ```
            # Add as monospace paragraph
            for cl in code_lines:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                run = p.add_run(cl)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
            continue

        # Table
        if line.startswith('|') and not line.startswith('| 算法'):
            # Check if next line is a separator row
            if i + 1 < len(lines) and re.match(r'^\|[\s\-:|]+\|$', lines[i + 1].rstrip()):
                header_row = line
                i += 2  # skip separator
                table_lines = [header_row]
                while i < len(lines) and lines[i].rstrip().startswith('|'):
                    table_lines.append(lines[i].rstrip())
                    i += 1

                # Parse markdown table
                rows = []
                for tl in table_lines:
                    cells = [c.strip() for c in tl.split('|')[1:-1]]
                    rows.append(cells)

                if len(rows) >= 2:
                    num_cols = len(rows[0])
                    table = doc.add_table(rows=len(rows), cols=num_cols, style='Table Grid')
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    for ri, row_data in enumerate(rows):
                        for ci, cell_text in enumerate(row_data):
                            cell = table.cell(ri, ci)
                            # Strip markdown bold
                            cell_text_clean = cell_text.replace('**', '')
                            p = cell.paragraphs[0]
                            p.paragraph_format.space_before = Pt(1)
                            p.paragraph_format.space_after = Pt(1)
                            run = p.add_run(cell_text_clean)
                            if ri == 0:
                                run.bold = True
                                run.font.size = Pt(10)
                            else:
                                run.font.size = Pt(10)
                            run.font.name = '宋体'
                            try:
                                rPr = run._element.get_or_add_rPr()
                                rFe = rPr.makeelement(qn('w:rFonts'), {})
                                rFe.set(qn('w:eastAsia'), '宋体')
                                rPr.insert(0, rFe)
                            except:
                                pass
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # Add spacing after table
                    doc.add_paragraph()
                continue

        # Blockquote
        if line.startswith('> '):
            text = line[2:].strip()
            # Also strip markdown bold: **text** -> text
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            run = p.add_run(text)
            run.italic = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # Ordered list
        if re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # strip bold for plain paragraph
            text = re.sub(r'`(.+?)`', r'\1', text)
            p = doc.add_paragraph(text, style='List Number')
            add_formatted_text(p, text)
            i += 1
            continue

        # Unordered list with bullet
        if line.startswith('- '):
            text = line[2:].strip()
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'`(.+?)`', r'\1', text)
            p = doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue

        # Regular paragraph
        if line.strip():
            text = line.strip()
            # Remove markdown bold markers for now (basic handling)
            text_clean = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text_clean = re.sub(r'`(.+?)`', r'\1', text_clean)
            text_clean = re.sub(r'\[(\d+)\]', r'[\1]', text_clean)  # keep citation brackets

            p = doc.add_paragraph()
            add_formatted_text(p, text_clean)
            i += 1
            continue

        i += 1

    # Add references section
    doc.add_heading('参考文献', level=1)
    refs = [
        "[1] Koklu, M., & Ozkan, I. A. (2020). Multiclass classification of dry beans using computer vision and machine learning techniques. Computers and Electronics in Agriculture, 174, 105507.",
        "[2] Chen, T., & Guestrin, C. (2016). XGBoost: A Scalable Tree Boosting System. Proceedings of the 22nd ACM SIGKDD, 785-794.",
        "[3] Pedregosa, F., et al. (2011). Scikit-learn: Machine Learning in Python. Journal of Machine Learning Research, 12, 2825-2830."
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.first_line_indent = Cm(-0.74)
        p.paragraph_format.left_indent = Cm(0.74)

    # Save
    doc.save('paper.docx')
    print("[OK] paper.docx generated successfully.")


def add_formatted_text(paragraph, text):
    """Add text with basic bold handling."""
    # Split by ** ** for bold segments
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


if __name__ == '__main__':
    create_word_from_md()
